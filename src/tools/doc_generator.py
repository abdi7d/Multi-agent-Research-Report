"""Document generator with graceful fallbacks when dependencies are missing.

This module will attempt to use `python-docx` and `reportlab` when available.
If they are not installed, it will produce plain text files so the pipeline
can complete during local development.
"""
import os
import json
from typing import Dict

# Try to import python-docx
try:
    from docx import Document  # type: ignore
    DOCX_AVAILABLE = True
except Exception:
    Document = None  # type: ignore
    DOCX_AVAILABLE = False

# Try to import reportlab
try:
    from reportlab.lib.pagesizes import letter  # type: ignore
    from reportlab.pdfgen import canvas  # type: ignore
    REPORTLAB_AVAILABLE = True
except Exception:
    letter = None  # type: ignore
    canvas = None  # type: ignore
    REPORTLAB_AVAILABLE = False


def _safe_text_from_report(report: Dict) -> str:
    parts = [report.get("title", "")]
    if report.get("summary"):
        parts.append(report.get("summary"))
    for sec in report.get("sections", []):
        heading = sec.get("heading", "")
        content = sec.get("content", "")
        parts.append(heading)
        # content may be list or string
        if isinstance(content, (list, tuple)):
            parts.extend([str(c) for c in content])
        else:
            parts.append(str(content))
    return "\n\n".join([p for p in parts if p])


def generate_docx(report: Dict, out_path: str):
    os.makedirs(os.path.dirname(out_path) or ".", exist_ok=True)
    if DOCX_AVAILABLE:
        doc = Document()
        doc.add_heading(report.get("title", ""), 0)
        if report.get("summary"):
            doc.add_paragraph(report.get("summary", ""))
        for sec in report.get("sections", []):
            doc.add_heading(sec.get("heading", "Section"), level=1)
            content = sec.get("content", "")
            if isinstance(content, (list, tuple)):
                for line in content:
                    doc.add_paragraph(str(line))
            else:
                doc.add_paragraph(str(content))
        doc.save(out_path)
        return out_path

    # Fallback: write a plain-text file with .docx extension so callers can still find it
    text = _safe_text_from_report(report)
    with open(out_path, "w", encoding="utf-8") as f:
        f.write(text)
    return out_path


def generate_pdf_from_text(report: Dict, out_path: str):
    os.makedirs(os.path.dirname(out_path) or ".", exist_ok=True)
    if REPORTLAB_AVAILABLE:
        c = canvas.Canvas(out_path, pagesize=letter)
        width, height = letter
        y = height - 50
        c.setFont("Helvetica-Bold", 16)
        c.drawString(50, y, report.get("title", ""))
        y -= 30
        if report.get("summary"):
            c.setFont("Helvetica", 11)
            for line in str(report.get("summary", "")).splitlines():
                c.drawString(50, y, line)
                y -= 14
                if y < 60:
                    c.showPage()
                    y = height - 50
        y -= 10
        for sec in report.get("sections", []):
            c.setFont("Helvetica-Bold", 12)
            c.drawString(50, y, sec.get("heading", ""))
            y -= 18
            c.setFont("Helvetica", 10)
            content = sec.get("content", "")
            if isinstance(content, (list, tuple)):
                lines = [str(l) for l in content]
            else:
                lines = str(content).splitlines()
            for line in lines:
                if y < 60:
                    c.showPage()
                    y = height - 50
                c.drawString(50, y, line)
                y -= 12
        c.save()
        return out_path

    # Fallback: write plain text to the path so a user can still read output
    text = _safe_text_from_report(report)
    with open(out_path, "w", encoding="utf-8") as f:
        f.write(text)
    return out_path
